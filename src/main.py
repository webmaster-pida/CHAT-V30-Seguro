# /src/main.py

import stripe
import json
import asyncio
import io
import re
from datetime import datetime, timedelta, timezone
from fastapi import FastAPI, Request, Depends, HTTPException, status, Form, Response
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any

# 👇 NUEVAS IMPORTACIONES AÑADIDAS
from pydantic import BaseModel
from urllib.parse import urlparse, parse_qs
from firebase_admin import auth as firebase_auth

# Librerías para documentos
from docx import Document
from fpdf import FPDF

from src.config import settings, log
from src.models.chat_models import ChatRequest, ChatMessage

from src.modules import perplexity_client, gemini_client, rag_client, firestore_client
from src.core.prompts import PIDA_SYSTEM_PROMPT
from src.core.security import get_current_user

from google.cloud import firestore
from google import genai
from google.genai import types

# Inicializar cliente global para utilidades dentro de main.py
genai_client = genai.Client(vertexai=True, project=settings.GOOGLE_CLOUD_PROJECT, location=settings.GOOGLE_CLOUD_LOCATION)

# --- MODELOS DE PETICIÓN ---
class VerificationRequest(BaseModel):
    frontend_url: str
    display_name: str | None = "Investigador"  # <-- Añadimos este campo

# 👇 NUEVO: Modelo para el Teaser de la Landing Page
class TeaserRequest(BaseModel):
    prompt: str

# MAPA DE TRADUCCIÓN: ID de Stripe -> Nombre del Plan interno para que no se equivoque
STRIPE_PRICE_MAP = {
    # BÁSICO (Añadir aquí IDs de Básico Anual / MXN si existen)
    "price_1SqFQiGgaloBN5L8U60ywohe": "basico", 
    "price_1SqFSFGgaloBN5L8kxegWZqC": "basico", 
    "price_1SqFSFGgaloBN5L8BMBeRPqb": "basico", 
    "price_1SqFSyGgaloBN5L8rrwrtUau": "basico", 
    
    # AVANZADO (Añadir aquí IDs de Avanzado Anual / MXN si existen)
    "price_1SqFUvGgaloBN5L8xOBssn6E": "avanzado",
    "price_1SqFWJGgaloBN5L8VKhkzLRH": "avanzado",
    "price_1SqFWJGgaloBN5L8roECNay2": "avanzado",
    "price_1SqFWJGgaloBN5L8hKpEvd1v": "avanzado",

    # PREMIUM (Añadir aquí IDs de Premium Anual / MXN si existen)
    "price_1SqFXIGgaloBN5L8vaGyleDT": "premium",
    "price_1SqFadGgaloBN5L86iwNYm1c": "premium",
    "price_1SqFadGgaloBN5L8AwTUeTSd": "premium",
    "price_1SqFadGgaloBN5L8QFHXe1i9": "premium",
}

# --- LÍMITES DE CHAT (POR PREGUNTA) ---
CHAT_LIMITS = {
    "basico": settings.LIMIT_BASICO_CHAT_MONTHLY,      
    "avanzado": settings.LIMIT_AVANZADO_CHAT_MONTHLY,  
    "premium": settings.LIMIT_PREMIUM_CHAT_MONTHLY,    
    "vip": -1                                        # Ilimitado
}

# Inicializar Stripe
stripe.api_key = settings.STRIPE_SECRET_KEY

app = FastAPI(
    title="PIDA Backend API",
    description="API para el asistente jurídico PIDA, con persistencia en BD y autenticación."
)

# --- CONFIGURACIÓN CORS ---
origins = settings.ALLOWED_ORIGINS

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_origin_regex=r"https://pida-ai-v20--.*\.web\.app$|https://.*\.app\.github\.dev$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- CLIENTE FIRESTORE ASÍNCRONO ---
db = firestore.AsyncClient()

# --- UTILIDADES DE DOCUMENTOS ---
def generate_filename(title: str, extension: str) -> str:
    safe_title = re.sub(r'[^a-zA-Z0-9áéíóúÁÉÍÓÚñÑ ]', '', title[:40])
    safe_title = safe_title.strip().replace(' ', '_')
    if not safe_title: safe_title = "Chat_PIDA"
    timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
    return f"{safe_title}_{timestamp}.{extension}"

def sanitize_text_for_pdf(text: str) -> str:
    if not text: return ""
    text = text.replace('$', '').replace('^{a}', 'a.').replace('^{o}', 'o.')
    
    replacements = { "•": "-", "—": "-", "–": "-", "“": '"', "”": '"', "‘": "'", "’": "'", "…": "...", "\u2013": "-", "\u2014": "-", "\u2022": "-", "\uF0B7": "-" }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text.encode('latin1', 'replace').decode('latin-1')

def write_markdown_to_pdf(pdf, text):
    import re
    pdf.set_font("Arial", "", 11)
    
    try:
        effective_page_width = pdf.epw
    except AttributeError:
        effective_page_width = pdf.w - pdf.l_margin - pdf.r_margin
        
    text = text.replace('$', '').replace('^{a}', 'a.').replace('^{o}', 'o.')
    
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # --- 1. PROCESAMIENTO DE TABLAS MARKDOWN ---
        if line.startswith('|') and line.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            for r_idx, t_line in enumerate(table_lines):
                cols = [c.strip() for c in t_line.split('|')[1:-1]]
                
                if all(re.match(r'^:?-+:?$', c) for c in cols):
                    continue
                if not cols:
                    continue
                    
                col_width = effective_page_width / len(cols)
                
                def get_cell_height(w, txt, is_bold):
                    pdf.set_font("Arial", "B" if is_bold else "", 10)
                    try: margin = pdf.c_margin
                    except: margin = 1
                    usable_w = w - (2 * margin)
                    
                    lines_count = 0
                    for p in str(txt).split('\n'):
                        words = p.split(' ')
                        if not words or (len(words) == 1 and words[0] == ''):
                            lines_count += 1; continue
                        curr_line = ""
                        for word in words:
                            if pdf.get_string_width(curr_line + word + " ") > usable_w and curr_line:
                                lines_count += 1; curr_line = word + " "
                            else: 
                                curr_line += word + " "
                        lines_count += 1
                    return lines_count * 6
                
                is_header = (r_idx == 0)
                max_height = max([get_cell_height(col_width, c.replace('**', '').replace('<br>', '\n').replace('<br/>', '\n'), is_bold=(is_header or '**' in c)) for c in cols] + [6])
                
                try: pb_trigger = pdf.page_break_trigger
                except: pb_trigger = pdf.h - pdf.b_margin
                if pdf.get_y() + max_height > pb_trigger:
                    pdf.add_page()
                    
                x_start = pdf.get_x()
                y_start = pdf.get_y()
                
                for c_idx, col in enumerate(cols):
                    col_clean = col.replace('**', '')
                    col_clean = re.sub(r'<br\s*/?>', '\n', col_clean, flags=re.IGNORECASE)
                    
                    if is_header:
                        pdf.set_fill_color(241, 245, 249)
                        pdf.rect(x_start + (c_idx * col_width), y_start, col_width, max_height, 'DF')
                    else:
                        pdf.rect(x_start + (c_idx * col_width), y_start, col_width, max_height)
                    
                    pdf.set_xy(x_start + (c_idx * col_width), y_start)
                    pdf.set_font("Arial", "B" if "装" in col or is_header else "", 10)
                    
                    if is_header:
                        pdf.set_text_color(29, 53, 87)
                    else:
                        pdf.set_text_color(0, 0, 0)
                        
                    pdf.multi_cell(col_width, 6, col_clean, border=0, align='L')
                
                pdf.set_xy(x_start, y_start + max_height)
            
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(5)
            continue

        if not line:
            pdf.ln(5)
            i += 1
            continue

        if line.startswith('## '):
            pdf.ln(3)
            pdf.set_font("Arial", "B", 13)
            pdf.set_text_color(29, 53, 87)
            clean_line = re.sub(r'<br\s*/?>', ' ', line.replace('## ', ''), flags=re.IGNORECASE)
            pdf.multi_cell(0, 8, clean_line)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)
            i += 1
            continue
            
        if line.startswith('### '):
            pdf.ln(2)
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(40, 70, 100)
            clean_line = re.sub(r'<br\s*/?>', ' ', line.replace('### ', ''), flags=re.IGNORECASE)
            pdf.multi_cell(0, 7, clean_line)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)
            i += 1
            continue

        if line.startswith('* ') or line.startswith('- '):
            pdf.set_x(15)
            line = "- " + line[2:]
        else:
            pdf.set_x(10)

        line = re.sub(r'<br\s*/?>', '', line, flags=re.IGNORECASE)

        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                pdf.set_font("Arial", "B", 11)
                pdf.write(6, part.strip('*'))
                pdf.set_font("Arial", "", 11)
            else:
                pdf.write(6, part)
        pdf.ln(6)
        
        i += 1

class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 14)
        self.set_text_color(29, 53, 87)
        self.cell(0, 10, "PIDA-AI: Historial de Chat", 0, 1, "L")
        self.set_font("Arial", "", 9)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Generado: {datetime.now().strftime('%d/%m/%Y, %H:%M:%S')}", 0, 1, "L")
        self.ln(5)
    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Pagina {self.page_no()}/{{nb}}", 0, 0, "C")

def create_chat_docx_sync(chat_text: str, title: str) -> tuple:
    from docx import Document
    from docx.shared import RGBColor
    import io
    import re
    from datetime import datetime

    doc = Document()
    
    heading = doc.add_heading(title, 0)
    heading.style.font.color.rgb = RGBColor(29, 53, 87)

    text = chat_text.replace('$', '').replace('^{a}', 'a.').replace('^{o}', 'o.')
    
    in_table = False
    current_table = None

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            in_table = False
            continue
            
        if line.startswith('|') and line.endswith('|'):
            if re.match(r'^\|?[\s\-:]+\|[\s\-:|]+$', line):
                continue
                
            cols = [c.strip() for c in line.strip('|').split('|')]
            
            if not in_table:
                current_table = doc.add_table(rows=1, cols=len(cols))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, col in enumerate(cols):
                    if i < len(hdr_cells):
                        cell_text = re.sub(r'<br\s*/?>', '\n', col.replace('**', ''), flags=re.IGNORECASE)
                        p = hdr_cells[i].paragraphs[0]
                        run = p.add_run(cell_text)
                        run.bold = True
                in_table = True
            else:
                if current_table:
                    row_cells = current_table.add_row().cells
                    for i, col in enumerate(cols):
                        if i < len(row_cells):
                            cell_text = re.sub(r'<br\s*/?>', '\n', col.replace('**', ''), flags=re.IGNORECASE)
                            row_cells[i].text = cell_text
            continue
        else:
            in_table = False

        line = re.sub(r'<br\s*/?>', '', line, flags=re.IGNORECASE)

        if line.startswith('## '):
            h = doc.add_heading(line.replace('## ', ''), level=2)
            h.style.font.color.rgb = RGBColor(29, 53, 87)
            continue
        
        if line.startswith('### '):
            h = doc.add_heading(line.replace('### ', ''), level=3)
            h.style.font.color.rgb = RGBColor(40, 70, 100)
            continue
        
        if line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            line = line[2:]
        else:
            p = doc.add_paragraph()

        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part.strip('*'))
                run.bold = True
            else:
                p.add_run(part)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Experto_PIDA_{timestamp}.docx"
    
    return doc_io.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename

def create_chat_pdf_sync(chat_text: str, title: str) -> tuple[bytes, str, str]:
    safe_text = sanitize_text_for_pdf(chat_text)
    safe_title = sanitize_text_for_pdf(title)
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font("Arial", "B", 12)
    pdf.multi_cell(0, 6, f"Tema: {safe_title}")
    pdf.ln(5)
    if not safe_text.strip():
        pdf.multi_cell(0, 6, "[Chat vacío]")
    else:
        write_markdown_to_pdf(pdf, safe_text)
    try:
        pdf_string = pdf.output(dest='S')
        if isinstance(pdf_string, str): pdf_bytes = pdf_string.encode('latin-1', 'replace')
        else: pdf_bytes = pdf_string
        stream = io.BytesIO(pdf_bytes)
        fname = generate_filename(title, "pdf")
        return stream.read(), "application/pdf", fname
    except Exception as e:
        err = FPDF()
        err.add_page()
        err.multi_cell(0, 10, f"Error: {str(e)}")
        return err.output(dest='S').encode('latin-1'), "application/pdf", "Error.pdf"

# --- VERIFICACIÓN DE SUSCRIPCIÓN ---
async def verify_active_subscription(current_user: Dict[str, Any]):
    user_id = current_user.get("uid")
    user_email = current_user.get("email", "").strip().lower()
    
    admin_domains = settings.ADMIN_DOMAINS
    admin_emails = settings.ADMIN_EMAILS
    email_domain = user_email.split("@")[-1] if "@" in user_email else ""

    # CONDICIÓN EXTRAORDINARIA VIP: Los bypasses de administración no se tocan
    if (email_domain in admin_domains) or (user_email in admin_emails):
        return

    try:
        user_doc = await db.collection("customers").document(user_id).get()
        if user_doc.exists and user_doc.to_dict().get("status") == "active":
            return 
        
        raise HTTPException(status_code=403, detail="Suscripción inactiva o requiere tarjeta válida.")
    except HTTPException as e: 
        raise e
    except Exception as e:
        log.error(f"Error Verificación: {e}")
        raise HTTPException(status_code=500, detail="Error de servidor.")

def get_date_utc_minus_6() -> str:
    utc_now = datetime.now(timezone.utc)
    cst_now = utc_now - timedelta(hours=6)
    return cst_now.strftime('%Y-%m-%d')

# NUEVA FUNCIÓN PARA LÍMITES MENSUALES
def get_month_utc_minus_6() -> str:
    utc_now = datetime.now(timezone.utc)
    cst_now = utc_now - timedelta(hours=6)
    return cst_now.strftime('%Y-%m') # Formato: YYYY-MM

# --- LÓGICA DE CONTROL DE LÍMITES POR PREGUNTA ---
async def consume_chat_credit(user_id: str, plan: str):
    plan_key = plan.lower().replace('á', 'a').strip()
    limit = CHAT_LIMITS.get(plan_key, 0)
    
    if limit == -1: return

    # Usamos el mes actual para agrupar los créditos de todo el mes
    current_month = get_month_utc_minus_6()
    stats_ref = db.collection('users').document(user_id).collection('usage_stats').document(current_month)
    
    @firestore.async_transactional
    async def check_and_increment(transaction, ref):
        snapshot = await ref.get(transaction=transaction)
        data = snapshot.to_dict() if snapshot.exists else {}
        current_count = data.get('chat_count', 0)
        
        if current_count >= limit:
            raise HTTPException(
                status_code=429, 
                detail=f"Límite mensual alcanzado para el plan {plan_key}" # Mensaje actualizado
            )
        
        transaction.set(ref, {
            'chat_count': current_count + 1,
            'last_updated': firestore.SERVER_TIMESTAMP
        }, merge=True)

    transaction = db.transaction()
    await check_and_increment(transaction, stats_ref)

async def refund_chat_credit(user_id: str):
    # Usamos el mes actual para devolver el crédito al mes correspondiente
    current_month = get_month_utc_minus_6()
    stats_ref = db.collection('users').document(user_id).collection('usage_stats').document(current_month)
    
    @firestore.async_transactional
    async def check_and_decrement(transaction, ref):
        snapshot = await ref.get(transaction=transaction)
        if snapshot.exists:
            data = snapshot.to_dict() or {}
            current_count = data.get('chat_count', 0)
            if current_count > 0:
                transaction.update(ref, {
                    'chat_count': current_count - 1,
                    'last_updated': firestore.SERVER_TIMESTAMP
                })
    
    try:
        transaction = db.transaction()
        await check_and_decrement(transaction, stats_ref)
    except Exception as e:
        log.error(f"Error procesando reembolso de crédito para {user_id}: {e}")

# --- GENERADOR STREAMING (EL ORQUESTADOR MÁGICO) ---
async def stream_chat_response_generator(chat_request: ChatRequest, country_code: str | None, user: Dict[str, Any], convo_id: str):
    user_id = user['uid']
    try:
        await verify_active_subscription(user) 
    except HTTPException as e:
        yield f"data: {json.dumps({'error': e.detail})}\n\n"
        return
    
    def create_sse_event(data: dict) -> str:
        return f"data: {json.dumps(data)}\n\n"

    try:
        history_from_db = await firestore_client.get_conversation_messages(user_id, convo_id)
        
        user_message = ChatMessage(role="user", content=chat_request.prompt)
        asyncio.create_task(firestore_client.add_message_to_conversation(user_id, convo_id, user_message))
        
        yield create_sse_event({"event": "status", "message": "Iniciando..."})
        
        history_for_gemini = gemini_client.prepare_history_for_genai(history_from_db)
        
        search_query = chat_request.prompt
        if history_from_db:
            yield create_sse_event({"event": "status", "message": "Contextualizando la búsqueda..."})
            try:
                recent_history = history_from_db[-4:] 
                context_text = "\n".join([f"{msg.role.upper()}: {msg.content}" for msg in recent_history])
                
                reformulation_prompt = f"""
Instrucción: Actúa como un clasificador binario y reformulador estricto. Tu ÚNICA tarea es decidir si la pregunta necesita buscarse en internet.

Historial reciente:
{context_text}

Pregunta actual: {chat_request.prompt}

Regla 1: Si la pregunta busca aclarar algo del historial ("¿De qué país hablamos?", "¿A qué te refieres?"), escribe ÚNICAMENTE la palabra: SKIP_SEARCH
Regla 2: Si la pregunta pide leyes o datos nuevos, reformúlala incluyendo el país principal ({country_code or 'el del historial'}).

Respuesta (sin comillas, sin explicaciones):"""
                
                response = await genai_client.aio.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=reformulation_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.0,
                        max_output_tokens=20,
                        automatic_function_calling=types.AutomaticFunctionCallingConfig(
                            disable=True
                        )
                    )
                )
                
                if response.text:
                    search_query = response.text.strip().replace('"', '').replace("'", "").replace('*', '')
                    log.info(f"Query original: '{chat_request.prompt}' | Query reformulada: '{search_query}'")
            except Exception as e:
                log.warning(f"Error reformulando la query, usando la original. Detalle: {e}")
                search_query = chat_request.prompt

        rag_context = ""
        web_context = ""

        if "SKIP_SEARCH" not in search_query.upper():
            yield create_sse_event({"event": "status", "message": "Analizando fuentes y biblioteca privada..."})
            
            rag_task = rag_client.search_internal_documents(search_query)
            perp_task = perplexity_client.get_perplexity_research(search_query)
            
            rag_context, web_context = await asyncio.gather(rag_task, perp_task)
            
            yield create_sse_event({"event": "status", "message": "Sintetizando y correlacionando fuentes..."})
        
        combined_context = f"{rag_context}\n{web_context}"
        trusted_urls_set = set(re.findall(r'https?://[^\s\)\],>]+', combined_context))
        
        yield create_sse_event({"event": "status", "message": "Formulando respuesta jurídica final..."})
        
        # 1. Se obtiene la fecha actual usando función existente
        fecha_actual = get_date_utc_minus_6()
        
        # 2. Inyecta la fecha en la primera línea del final_prompt
        final_prompt = f"""Fecha actual del sistema: {fecha_actual}
Contexto geográfico principal: {country_code or 'General'}

Toma en cuenta las fuentes proporcionadas. 
IMPORTANTE: No uses '[INVESTIGACIÓN WEB RECIENTE]' como nombre de fuente. Extrae el nombre real del sitio web (ej: ONU, Amnistía, Wikipedia) desde la URL proporcionada.

[CONTEXTO INTERNO DE JURISPRUDENCIA (RAG)]
(⚠️ REGLA ESTRICTA: Tienes PROHIBIDO extraer o mostrar URLs de este bloque. Usa solo el conocimiento en texto plano).
{rag_context}

[INVESTIGACIÓN WEB RECIENTE (Perplexity)]
(✅ REGLA ESTRICTA Y FILTRO GEOGRÁFICO: Debes usar las URLs de este bloque y convertirlas en hipervínculos Markdown. **EXCEPCIÓN CRÍTICA:** Si tu 'Contexto geográfico principal' es un país (ej. El Salvador) y la investigación web te trae leyes o instituciones de OTRO PAÍS distinto (ej. el BOE de España, Congreso de España), **TIENES ESTRICTAMENTE PROHIBIDO** usar y citar esas fuentes extranjeras. Limítate a usar fuentes del país correcto, doctrina general o de organismos internacionales).
{web_context}

---
INSTRUCCIÓN CRÍTICA DE ENLACES: 
1. ¡NO USES NÚMEROS ENTRE CORCHETES COMO [1] O [2] PARA CITAR! El sistema los borrará automáticamente y perderemos la referencia.
2. Tienes que leer la sección "FUENTES DE INTERNET" que te dio Perplexity y crear hipervínculos Markdown reales (ej: [Nombre de la Institución](URL_COMPLETA)).
3. ES OBLIGATORIO que los enlaces válidos aparezcan incrustados dentro de los párrafos. Si todas las fuentes web fueron descartadas por ser de otro país irrelevante, básate solo en tu conocimiento y el RAG, y no pongas enlaces web.

Pregunta del usuario: {chat_request.prompt}
⚠️ REGLA FINAL: Verifica la existencia real de lo que pide el usuario antes de responder. No asumas su premisa como verdadera.
"""
        
        full_response_text = ""
        
        async for chunk in gemini_client.generate_streaming_response(
            system_prompt=PIDA_SYSTEM_PROMPT,
            prompt=final_prompt,
            history=history_for_gemini,
            trusted_urls=trusted_urls_set 
        ):
            yield create_sse_event({'text': chunk})
            full_response_text += chunk

        if full_response_text:
            model_message = ChatMessage(role="model", content=full_response_text)
            await firestore_client.add_message_to_conversation(user_id, convo_id, model_message)

        yield create_sse_event({'event': 'done'})

    except Exception as e:
        log.error(f"Error crítico streaming convo {convo_id}: {e}", exc_info=True)
        error_message = json.dumps({"error": "Ocurrió un error interno al generar la respuesta."})
        yield f"data: {error_message}\n\n"

# --- ENDPOINTS ---

@app.get("/status", tags=["Status"])
def read_status():
    return {"status": "ok", "message": "PIDA Chat Backend v3.0 (Security Patched)"}

# 👇 ENDPOINT PARA EL LEAD MAGNET (TRY-BEFORE-YOU-BUY) 100% REAL Y COMPLETO
@app.post("/teaser-chat", tags=["Lead Magnet"])
async def teaser_chat_stream_handler(request: Request, body: TeaserRequest):
    """
    Endpoint público para la Landing Page.
    Genera una respuesta idéntica a la de la app (Perplexity + RAG + Respuesta completa).
    Protegido por ID de Navegador (Max 2 al día) para evitar abusos de consumo.
    """
    country_code = request.headers.get('X-Country-Code', 'General')
    
    # 1. OBTENER EL SELLO DEL NAVEGADOR
    anon_id = request.headers.get("X-Anon-ID")
    
    if not anon_id:
        raise HTTPException(status_code=400, detail="Falta identificador de dispositivo.")
    
    # 2. VALIDAR LÍMITES EN FIRESTORE POR NAVEGADOR
    today_str = datetime.now(timezone.utc).strftime('%Y-%m-%d')
    limit_doc_ref = db.collection('teaser_limits').document(f"{anon_id}_{today_str}")
    
    try:
        doc = await limit_doc_ref.get()
        if doc.exists:
            count = doc.to_dict().get('count', 0)
            if count >= 2:  # 👈 LÍMITE: 2 pruebas completas por navegador al día
                raise HTTPException(status_code=429, detail="Has alcanzado el límite de demostraciones gratuitas por hoy.")
            await limit_doc_ref.update({'count': firestore.Increment(1)})
        else:
            await limit_doc_ref.set({'count': 1})
    except HTTPException as he:
        raise he
    except Exception as e:
        log.error(f"Error verificando límite de Teaser para {anon_id}: {e}")
    
    async def restricted_stream_generator():
        try:
            # 3. BÚSQUEDA 100% REAL EN PARALELO (RAG + PERPLEXITY)
            yield f"data: {json.dumps({'event': 'status', 'message': 'Analizando fuentes y biblioteca privada...'})}\n\n"
            
            search_query = body.prompt
            
            rag_task = rag_client.search_internal_documents(search_query)
            perp_task = perplexity_client.get_perplexity_research(search_query)
            
            rag_context, web_context = await asyncio.gather(rag_task, perp_task)
            
            yield f"data: {json.dumps({'event': 'status', 'message': 'Sintetizando y correlacionando fuentes...'})}\n\n"
            
            combined_context = f"{rag_context}\n{web_context}"
            trusted_urls_set = set(re.findall(r'https?://[^\s\)\],>]+', combined_context))
            
            yield f"data: {json.dumps({'event': 'status', 'message': 'Formulando respuesta jurídica final...'})}\n\n"
            
            fecha_actual = get_date_utc_minus_6()
            
            # EL MISMO PROMPT EXACTO DE PRODUCCIÓN DE TU CHAT PREMIUM
            final_prompt = f"""Fecha actual del sistema: {fecha_actual}
Contexto geográfico principal: {country_code or 'General'}

Toma en cuenta las fuentes proporcionadas. 
IMPORTANTE: No uses '[INVESTIGACIÓN WEB RECIENTE]' como nombre de fuente. Extrae el nombre real del sitio web (ej: ONU, Amnistía, Wikipedia) desde la URL proporcionada.

[CONTEXTO INTERNO DE JURISPRUDENCIA (RAG)]
(⚠️ REGLA ESTRICTA: Tienes PROHIBIDO extraer o mostrar URLs de este bloque. Usa solo el conocimiento en texto plano).
{rag_context}

[INVESTIGACIÓN WEB RECIENTE (Perplexity)]
(✅ REGLA ESTRICTA Y FILTRO GEOGRÁFICO: Debes usar las URLs de este bloque y convertirlas en hipervínculos Markdown. **EXCEPCIÓN CRÍTICA:** Si tu 'Contexto geográfico principal' es un país (ej. El Salvador) y la investigación web te trae leyes o instituciones de OTRO PAÍS distinto (ej. el BOE de España, Congreso de España), **TIENES ESTRICTAMENTE PROHIBIDO** usar y citar esas fuentes extranjeras. Limítate a usar fuentes del país correcto, doctrina general o de organismos internacionales).
{web_context}

---
INSTRUCCIÓN CRÍTICA DE ENLACES: 
1. ¡NO USES NÚMEROS ENTRE CORCHETES COMO [1] O [2] PARA CITAR! El sistema los borrará automáticamente y perderemos la referencia.
2. Tienes que leer la sección "FUENTES DE INTERNET" que te dio Perplexity y crear hipervínculos Markdown reales (ej: [Nombre de la Institución](URL_COMPLETA)).
3. ES OBLIGATORIO que los enlaces válidos aparezcan incrustados dentro de los párrafos. Si todas las fuentes web fueron descartadas por ser de otro país irrelevante, básate solo en tu conocimiento y el RAG, y no pongas enlaces web.

Pregunta del usuario: {body.prompt}
⚠️ REGLA FINAL: Verifica la existencia real de lo que pide el usuario antes de responder. No asumas su premisa como verdadera.
"""
            # 👇 RESPUESTA COMPLETA, SIN CORTES (Genera todo, incluyendo <pida_questions>)
            async for chunk in gemini_client.generate_streaming_response(
                system_prompt=PIDA_SYSTEM_PROMPT,
                prompt=final_prompt,
                history=[],
                trusted_urls=trusted_urls_set 
            ):
                yield f"data: {json.dumps({'text': chunk})}\n\n"
                
            yield f"data: {json.dumps({'event': 'done'})}\n\n"
            
        except Exception as e:
            log.error(f"Error en teaser chat: {e}", exc_info=True)
            yield f"data: {json.dumps({'error': 'Error procesando tu consulta inicial.'})}\n\n"

    headers = { 
        "Content-Type": "text/event-stream", 
        "Cache-Control": "no-cache", 
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no"
    }
    return StreamingResponse(restricted_stream_generator(), headers=headers)

# 👇 NUEVO ENDPOINT DE VERIFICACIÓN DINÁMICO AÑADIDO
@app.post("/send-verification-email", tags=["Security"])
async def send_custom_verification(
    payload: VerificationRequest, 
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    try:
        user_email = current_user.get("email")
        
        # 1. Configurar el retorno base de Firebase
        action_code_settings = firebase_auth.ActionCodeSettings(
            url=payload.frontend_url,
            handle_code_in_app=False
        )
        
        # 2. Generar el link estándar interno de Firebase
        firebase_link = firebase_auth.generate_email_verification_link(
            user_email, 
            action_code_settings
        )
        
        # 3. Extraer el oobCode (Token de seguridad) de forma quirúrgica
        parsed_url = urlparse(firebase_link)
        queries = parse_qs(parsed_url.query)
        oob_code = queries.get("oobCode")[0]
        
        # 4. Construir tu link limpio hacia tu ruta de React (Evita pantallas de Google)
        custom_clean_link = f"{payload.frontend_url}/auth-action?mode=verifyEmail&oobCode={oob_code}"
        
        # 5. Despachar usando tu propio cliente de notificaciones
        await firestore_client.send_email_notification(
            to_email=user_email,
            template_name='email-verification',
            template_data={
                'verificationLink': custom_clean_link, 
                'displayName': payload.display_name
            }
        )
        
        log.info(f"Link de verificación dinámico generado para {user_email}")
        return {"status": "ok", "message": "Enlace de verificación enviado correctamente."}
        
    except Exception as e:
        log.error(f"Error generando correo de verificación: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Error interno al generar el enlace de seguridad.")

@app.get("/conversations", response_model=List[Dict[str, Any]], tags=["Chat History"])
async def get_user_conversations(current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    return await firestore_client.get_conversations(current_user['uid'])

@app.get("/conversations/{convo_id}/messages", response_model=List[ChatMessage], tags=["Chat History"])
async def get_conversation_details(convo_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    return await firestore_client.get_conversation_messages(current_user['uid'], convo_id)

@app.post("/conversations", response_model=Dict[str, Any], status_code=status.HTTP_201_CREATED, tags=["Chat History"])
async def create_new_empty_conversation(request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    body = await request.json()
    title = body.get("title", "Nuevo Chat")
    if not title: raise HTTPException(400, "El título no puede estar vacío")
    new_convo = await firestore_client.create_new_conversation(current_user['uid'], title)
    return new_convo

@app.delete("/conversations/{convo_id}", status_code=status.HTTP_204_NO_CONTENT, tags=["Chat History"])
async def delete_a_conversation(convo_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    await firestore_client.delete_conversation(current_user['uid'], convo_id)
    return

@app.patch("/conversations/{convo_id}/title", status_code=status.HTTP_204_NO_CONTENT, tags=["Chat History"])
async def update_conversation_title_handler(convo_id: str, request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    await verify_active_subscription(current_user)
    body = await request.json()
    new_title = body.get("title")
    if not new_title: raise HTTPException(400, "El título no puede estar vacío")
    await firestore_client.update_conversation_title(current_user['uid'], convo_id, new_title)
    return

@app.post("/chat-stream/{convo_id}", tags=["Chat"])
async def chat_stream_handler(
    convo_id: str, 
    chat_request: ChatRequest, 
    request: Request, 
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    country_code = request.headers.get('X-Country-Code', None)
    user_id = current_user['uid']
    user_email = current_user.get('email', '').strip().lower()
    user_plan = 'none' 

    admin_domains = settings.ADMIN_DOMAINS
    admin_emails = settings.ADMIN_EMAILS
    email_domain = user_email.split("@")[-1] if "@" in user_email else ""

    if (email_domain in admin_domains) or (user_email in admin_emails):
        user_plan = 'vip'
    else:
        try:
            cust_doc = await db.collection('customers').document(user_id).get()
            if cust_doc.exists:
                data = cust_doc.to_dict()
                if data.get('status') == 'active':
                    user_plan = data.get('plan', 'basico')
        except Exception as e:
            log.error(f"Error obteniendo plan usuario: {e}")

    await consume_chat_credit(user_id, user_plan)

    async def counted_stream_generator():
        has_error = False
        tokens_sent = False
        
        try:
            async for chunk in stream_chat_response_generator(
                chat_request, 
                country_code, 
                current_user, 
                convo_id
            ):
                if '"error":' in chunk:
                    has_error = True
                
                if '"text":' in chunk and not has_error:
                    tokens_sent = True
                    
                yield chunk
        finally:
            if has_error or not tokens_sent:
                asyncio.create_task(refund_chat_credit(user_id))

    headers = { 
        "Content-Type": "text/event-stream", 
        "Cache-Control": "no-cache", 
        "Connection": "keep-alive", 
        "X-Accel-Buffering": "no" 
    }
    
    return StreamingResponse(counted_stream_generator(), headers=headers)

@app.post("/download-chat", tags=["Chat"])
async def download_chat(
    convo_id: str = Form(...),
    file_format: str = Form("docx"),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    try:
        user_id = current_user['uid']
        
        messages = await firestore_client.get_conversation_messages(user_id, convo_id)
        if not messages:
            raise HTTPException(status_code=404, detail="Conversación no encontrada o vacía.")
            
        title = "Exportación de Chat PIDA"
        
        chat_lines = []
        for msg in messages:
            role_str = "Usuario" if msg.role == "user" else "PIDA"
            content = msg.content
            
            if msg.role == "model":
                content = content.replace("_Fin del análisis._", "")
                
                if "<pida_questions>" in content and "</pida_questions>" in content:
                    def replacer(match):
                        q_raw = match.group(1)
                        qs = [q.strip() for q in q_raw.split('|') if q.strip()]
                        if not qs: return ""
                        
                        res = "\n\n**Preguntas de seguimiento sugeridas:**\n"
                        for q in qs: 
                            res += f"- {q}\n"
                        return res
                    
                    content = re.sub(r"<pida_questions>(.*?)</pida_questions>", replacer, content, flags=re.DOTALL)

            chat_lines.append(f"**{role_str}:**\n{content.strip()}")
            
        chat_text = "\n\n".join(chat_lines)
        
        if len(chat_text) > settings.MAX_EXPORT_LENGTH:
            chat_text = chat_text[:settings.MAX_EXPORT_LENGTH] + "\n\n[Texto truncado por límite de seguridad]"

        if file_format.lower() == "docx":
            content_bytes, mime, fname = await asyncio.to_thread(create_chat_docx_sync, chat_text, title)
        else:
            content_bytes, mime, fname = await asyncio.to_thread(create_chat_pdf_sync, chat_text, title)
            
        return Response(content=content_bytes, media_type=mime, headers={"Content-Disposition": f"attachment; filename={fname}"})
        
    except HTTPException as he:
        raise he
    except Exception as e:
        log.error(f"Error descarga chat: {e}")
        raise HTTPException(500, f"Error generando archivo: {e}")


@app.post("/check-vip-access", tags=["Security"])
async def check_vip_access_handler(current_user: Dict[str, Any] = Depends(get_current_user)):
    user_email = current_user.get("email", "").strip().lower()
    admin_domains = settings.ADMIN_DOMAINS
    admin_emails = settings.ADMIN_EMAILS
    email_domain = user_email.split("@")[-1] if "@" in user_email else ""
    if (email_domain in admin_domains) or (user_email in admin_emails):
        return {"is_vip_user": True}
    return {"is_vip_user": False}

@app.post("/validate-promo-code", tags=["Billing"])
async def validate_promo_code(request: Request):
    try:
        data = await request.json()
        promo_code = data.get("code", "").strip().upper() 
        price_id = data.get("priceId")

        if not promo_code or not price_id:
            raise HTTPException(status_code=400, detail="Faltan datos requeridos.")

        current_plan_name = STRIPE_PRICE_MAP.get(price_id)
        # Se elimina la restricción inicial para no rechazar precios válidos no mapeados (ej. Anuales o MXN nuevos)
        # Esto garantiza la compatibilidad con cupones globales. Bloquearemos más adelante solo si el cupón tiene restricciones de plan.

        try:
            promos = stripe.PromotionCode.list(code=promo_code, active=True, limit=1)
        except stripe.error.StripeError as e:
            log.error(f"StripeError al buscar promoción: {e.user_message}")
            raise HTTPException(status_code=400, detail="Error de conexión con Stripe al buscar el código.")

        if not hasattr(promos, 'data') or not promos.data:
            raise HTTPException(status_code=404, detail="El código promocional no es válido o ha expirado.")

        promo_obj = promos.data[0]
        
        promo_coupon = None
        if hasattr(promo_obj, 'promotion') and promo_obj.promotion:
            promo_coupon = promo_obj.promotion.coupon
        elif hasattr(promo_obj, 'coupon'):
            promo_coupon = promo_obj.coupon
            
        if not promo_coupon:
            raise HTTPException(status_code=404, detail="No se encontró un cupón asociado a esta promoción.")

        if isinstance(promo_coupon, str):
            coupon_id = promo_coupon
        elif hasattr(promo_coupon, 'id'):
            coupon_id = promo_coupon.id
        elif isinstance(promo_coupon, dict):
            coupon_id = promo_coupon.get('id')
        else:
            coupon_id = str(promo_coupon)

        try:
            coupon = stripe.Coupon.retrieve(coupon_id)
            price_obj = stripe.Price.retrieve(price_id)
        except stripe.error.StripeError as e:
            error_msg = e.user_message or str(e)
            log.error(f"StripeError al recuperar cupón '{coupon_id}' o precio '{price_id}': {error_msg}")
            raise HTTPException(status_code=400, detail=f"No se pudo validar el cupón en Stripe: {error_msg}")

        if hasattr(coupon, 'metadata') and coupon.metadata and "allowed_plans" in coupon.metadata:
            allowed_plans_meta = coupon.metadata["allowed_plans"]
            allowed_list = [p.strip().lower() for p in allowed_plans_meta.split(",")]
            
            if not current_plan_name:
                raise HTTPException(
                    status_code=400, 
                    detail="No se pudo determinar el plan para este precio. Asegúrate de añadir este ID a STRIPE_PRICE_MAP para usar cupones específicos."
                )

            if current_plan_name not in allowed_list:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Este cupón solo es válido para el plan {allowed_plans_meta.upper()}."
                )
                
        elif hasattr(coupon, 'applies_to') and coupon.applies_to and hasattr(coupon.applies_to, 'products'):
            current_product_id = price_obj.product if hasattr(price_obj, 'product') else None
            allowed_products = coupon.applies_to.products
            
            if current_product_id not in allowed_products:
                raise HTTPException(status_code=400, detail="Código no válido para este plan.")

        original_amount = price_obj.unit_amount if hasattr(price_obj, 'unit_amount') else None
        if original_amount is None:
            raise HTTPException(status_code=400, detail="El precio no tiene un monto fijo compatible con descuentos.")
            
        currency = price_obj.currency.upper() if hasattr(price_obj, 'currency') and price_obj.currency else ""
        final_amount = original_amount
        discount_desc = ""

        percent_off = coupon.percent_off if hasattr(coupon, 'percent_off') else None
        amount_off = coupon.amount_off if hasattr(coupon, 'amount_off') else None
        coupon_currency = coupon.currency if hasattr(coupon, 'currency') else None

        if percent_off is not None:
            discount_amount = int(round(original_amount * (percent_off / 100)))
            final_amount = original_amount - discount_amount
            discount_desc = f"-{percent_off}%"
            
        elif amount_off is not None:
            if coupon_currency and coupon_currency.upper() != currency:
                raise HTTPException(status_code=400, detail="La moneda del cupón no coincide con la del plan.")
            final_amount = original_amount - amount_off
            discount_desc = f"-${amount_off / 100:.2f} {currency}"

        if final_amount < 0: 
            final_amount = 0

        coupon_name = coupon.name if hasattr(coupon, 'name') and coupon.name else promo_obj.code

        return {
            "valid": True,
            "code": promo_obj.code,
            "original_amount": original_amount,
            "final_amount": final_amount,
            "currency": currency,
            "description": discount_desc,
            "coupon_name": coupon_name,
            "promo_id": promo_obj.id
        }

    except HTTPException as he:
        raise he 
    except Exception as e:
        log.error(f"Error de sistema no controlado: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Error interno del servidor.")
        
@app.post("/create-payment-intent", tags=["Billing"])
async def create_payment_intent(data: Dict[str, Any], current_user: Dict[str, Any] = Depends(get_current_user)):
    user_email = current_user.get("email", "").strip().lower()
    email_domain = user_email.split("@")[-1] if "@" in user_email else ""
    
    # 1. Bypass inteligente para las cuentas VIP del backend en variables de entorno
    is_admin = (email_domain in settings.ADMIN_DOMAINS) or (user_email in settings.ADMIN_EMAILS)
    
    # 2. BLINDAJE DE STRIPE: Impedir intenciones de cobro a correos no verificados, a menos que sean admin
    # (Comentado temporalmente para permitir creación de suscripciones sin verificar email)
    # if not current_user.get("email_verified", False) and not is_admin:
    #     raise HTTPException(
    #         status_code=403, 
    #         detail="Operación denegada. Debes verificar tu correo electrónico antes de adquirir un plan."
    #     )

    try:
        user_email = current_user.get("email")
        uid = current_user["uid"]
        
        price_id = data.get("priceId")
        plan_key = STRIPE_PRICE_MAP.get(price_id) 
        if not plan_key:
            raise HTTPException(status_code=400, detail="ID de Precio no reconocido o alterado. Operación denegada.")
            
        customer_name = data.get("name", "") 
        user_promo_code = data.get("promotion_code", "").strip()
        
        payment_method_id = data.get("paymentMethodId")
        if not payment_method_id:
            raise HTTPException(status_code=400, detail="Es necesario un método de pago válido.")

        customer = None
        search_query = f"metadata['firebaseUID']:'{uid}' OR metadata['uid']:'{uid}'"
        search_result = stripe.Customer.search(query=search_query, limit=1)
        if search_result.data:
            customer = search_result.data[0]
        else:
            existing_customers = stripe.Customer.list(email=user_email, limit=1)
            if existing_customers.data:
                customer = existing_customers.data[0]

        trial_days = 5
        trial_historico_usado = False
        try:
            if customer:
                customer_doc = await db.collection("customers").document(uid).get()
                if customer_doc.exists and customer_doc.to_dict().get("trial_used", False) is True:
                    trial_historico_usado = True
        except Exception as e:
            log.error(f"Error verificando trial histórico: {e}")

        if trial_historico_usado:
            trial_days = 0 
            
        pm = stripe.PaymentMethod.retrieve(payment_method_id)
        pm_fingerprint = pm.card.fingerprint if pm.type == 'card' else None

        if trial_days > 0 and pm_fingerprint:
            existing_user_query = db.collection("customers").where("email", "==", user_email).where("trial_used", "==", True).limit(1)
            docs = await existing_user_query.get()
            if len(docs) > 0:
                trial_days = 0

        if not customer:
            customer = stripe.Customer.create(
                email=user_email, 
                name=customer_name, 
                metadata={"uid": uid, "firebaseUID": uid}
            )
        else:
            if customer_name and customer.name != customer_name:
                stripe.Customer.modify(customer.id, name=customer_name)

        promo_id = None
        if user_promo_code:
            promos = stripe.PromotionCode.list(code=user_promo_code, active=True, limit=1)
            if promos.data: 
                promo_obj = promos.data[0]
                coupon_id = promo_obj.coupon.id
                try:
                    coupon = stripe.Coupon.retrieve(coupon_id)
                    allowed_plans_meta = coupon.metadata.get("allowed_plans")
                    if allowed_plans_meta:
                        allowed_list = [p.strip().lower() for p in allowed_plans_meta.split(",")]
                        if plan_key.lower() not in allowed_list:
                            raise HTTPException(status_code=400, detail=f"Cupón inválido para el plan {plan_key}.")
                    elif coupon.get("applies_to"):
                        price_obj = stripe.Price.retrieve(price_id)
                        current_product_id = price_obj.product
                        allowed_products = coupon.applies_to.get("products", [])
                        if allowed_products and current_product_id not in allowed_products:
                            raise HTTPException(status_code=400, detail="El código no es válido para este nivel de plan.")
                except Exception as e:
                    if isinstance(e, HTTPException): raise e
                    raise HTTPException(status_code=400, detail="Error al validar restricciones del producto en Stripe.")
                promo_id = promo_obj.id
            else:
                raise HTTPException(status_code=400, detail=f"Código promocional inválido o expirado.")

        stripe.PaymentMethod.attach(payment_method_id, customer=customer.id)
        stripe.Customer.modify(
            customer.id,
            invoice_settings={"default_payment_method": payment_method_id}
        )
        
        existing_subs = stripe.Subscription.list(customer=customer.id, limit=1)

        if existing_subs.data:
            sub = existing_subs.data[0]
            sub_item_id = sub['items']['data'][0].id
            
            modify_params = {
                "default_payment_method": payment_method_id,
                "expand": ['latest_invoice.payment_intent', 'pending_setup_intent'],
                "metadata": {"uid": uid, "plan_key": plan_key}
            }
            
            if promo_id: 
                modify_params["promotion_code"] = promo_id
                
            if sub['items']['data'][0].price.id != price_id:
                modify_params["items"] = [{"id": sub_item_id, "price": price_id}]
                
            subscription = stripe.Subscription.modify(sub.id, **modify_params)
            
            if subscription.status in ['past_due', 'incomplete'] and subscription.latest_invoice:
                try:
                    invoice_id = subscription.latest_invoice.id if hasattr(subscription.latest_invoice, 'id') else subscription.latest_invoice
                    stripe.Invoice.pay(invoice_id)
                    subscription = stripe.Subscription.retrieve(sub.id, expand=['latest_invoice.payment_intent', 'pending_setup_intent'])
                except stripe.error.CardError as e:
                    raise HTTPException(status_code=400, detail=f"La nueva tarjeta también fue declinada: {e.user_message}")
                except Exception as e:
                    log.error(f"Error al procesar el pago pendiente: {e}")

        else:
            subscription = stripe.Subscription.create(
                customer=customer.id,
                items=[{'price': price_id}],
                trial_period_days=trial_days if trial_days > 0 else None,
                promotion_code=promo_id, 
                default_payment_method=payment_method_id, 
                expand=['latest_invoice.payment_intent', 'pending_setup_intent'], 
                metadata={"uid": uid, "plan_key": plan_key}
            )

        if subscription.status == 'incomplete' and subscription.latest_invoice and subscription.latest_invoice.payment_intent:
            return {"clientSecret": subscription.latest_invoice.payment_intent.client_secret, "requiresAction": True}
        
        if subscription.status == 'trialing' and subscription.pending_setup_intent:
             return {"clientSecret": subscription.pending_setup_intent.client_secret, "requiresAction": True}
            
        return {"subscriptionId": subscription.id, "success": True, "requiresAction": False}

    except HTTPException as he:
        raise he
    except Exception as e:
        log.error(f"Error Suscripción: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/stripe-webhook", tags=["Billing"])
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig_header = request.headers.get("Stripe-Signature")
    
    try:
        webhook_secret = settings.STRIPE_WEBHOOK_SECRET 
        if not webhook_secret:
            log.error("⚠️ STRIPE_WEBHOOK_SECRET no está configurado.")
            return Response(content="Webhook secret missing", status_code=500)

        stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
        
        event_dict = json.loads(payload)
        event_type = event_dict.get('type')
        data_object = event_dict.get('data', {}).get('object', {})
        
        log.info(f"📩 Webhook recibido de forma segura: {event_type}")

        def resolve_plan(sub_obj):
            try:
                items = sub_obj.get('items', {})
                if not items: return "none"
                
                data_list = items.get('data', [])
                if not data_list: return "none"
                
                p_id = data_list[0].get('price', {}).get('id')
                return STRIPE_PRICE_MAP.get(p_id, "none")
            except Exception as e:
                log.error(f"Error en resolve_plan: {e}")
                return "none"

        if event_type in ['customer.subscription.created', 'customer.subscription.updated']:
            metadata = data_object.get('metadata') or {}  
            uid = metadata.get('uid')
            stripe_status = data_object.get('status')
            
            customer_id = data_object.get('customer')
            customer_email = None
            if customer_id:
                try:
                    stripe_cust = stripe.Customer.retrieve(customer_id)
                    customer_email = getattr(stripe_cust, 'email', None)
                except Exception as e:
                    log.error(f"Error recuperando email del cliente {customer_id}: {e}")

            if uid:
                is_active = stripe_status in ['active', 'trialing']
                is_trial = (stripe_status == 'trialing') 
                
                update_data = {
                    "status": "active" if is_active else "inactive",
                    "plan": resolve_plan(data_object) if is_active else "none",
                    "stripe_status": stripe_status,
                    "has_trial": is_trial,
                    "updated_at": firestore.SERVER_TIMESTAMP
                }
                
                if customer_email:
                    update_data["email"] = customer_email
                
                if is_trial:
                    update_data["trial_used"] = True

                await db.collection("customers").document(uid).set(update_data, merge=True)
                log.info(f"🛡️ Webhook: {uid} set to {'active' if is_active else 'inactive'} ({stripe_status})")

                if is_active and event_type == 'customer.subscription.created':
                    try:
                        user_doc = await db.collection('users').document(uid).get()
                        user_data = user_doc.to_dict() if user_doc.exists else {}
                        
                        customer_email = user_data.get('email')
                        customer_name = user_data.get('displayName', 'Investigador')

                        if not customer_email:
                            customer_id = data_object.get('customer')
                            if customer_id:
                                stripe_cust = stripe.Customer.retrieve(customer_id)
                                customer_email = getattr(stripe_cust, 'email', None)
                                stripe_name = getattr(stripe_cust, 'name', None)
                                
                                if not customer_name or customer_name == 'Investigador':
                                    customer_name = stripe_name if stripe_name else 'Investigador'

                        if customer_email:
                            await firestore_client.send_email_notification(
                                to_email=customer_email,
                                template_name='welcome-trial',
                                template_data={'displayName': customer_name}
                            )

                            await firestore_client.send_email_notification(
                                to_email="contacto@pida-ai.com",
                                template_name='admin-notification',
                                template_data={
                                    'customerName': customer_name,
                                    'customerEmail': customer_email,
                                    'planName': resolve_plan(data_object),
                                    'date': datetime.now().strftime("%d/%m/%Y %H:%M")
                                }
                            )
                        else:
                            log.warning(f"⚠️ No se pudo enviar correo. No se encontró el email para el uid: {uid}")

                    except Exception as e:
                        log.error(f"Error al intentar enviar correos: {e}")

        elif event_type == 'invoice.payment_succeeded':
            subscription_id = data_object.get('subscription')
            if subscription_id:
                try:
                    sub = stripe.Subscription.retrieve(subscription_id)
                    metadata = getattr(sub, 'metadata', {})
                    uid = metadata.get('uid')
                    
                    customer_id = data_object.get('customer')
                    customer_email = None
                    if customer_id:
                        stripe_cust = stripe.Customer.retrieve(customer_id)
                        customer_email = getattr(stripe_cust, 'email', None)

                    if uid:
                        update_data = {
                            "status": "active",
                            "updated_at": firestore.SERVER_TIMESTAMP
                        }
                        
                        if customer_email:
                            update_data["email"] = customer_email
                            
                        await db.collection("customers").document(uid).set(update_data, merge=True)
                        log.info(f"✅ Webhook: Pago exitoso para {uid}. Estado activado.")
                except Exception as e:
                    log.error(f"Error procesando invoice.payment_succeeded: {e}")

        elif event_type in ['customer.subscription.deleted', 'invoice.payment_failed']:
            metadata = data_object.get('metadata') or {}
            uid = metadata.get('uid')
            customer_id = data_object.get('customer')
            
            if not uid and data_object.get('subscription'):
                try:
                    sub = stripe.Subscription.retrieve(data_object.get('subscription'))
                    sub_meta = sub.get('metadata') or {}
                    uid = sub_meta.get('uid')
                    if not customer_id:
                        customer_id = sub.get('customer')
                except Exception: pass
            
            if uid and customer_id:
                try:
                    active_subs = stripe.Subscription.list(customer=customer_id, status='active', limit=1)
                    trial_subs = stripe.Subscription.list(customer=customer_id, status='trialing', limit=1)
                    
                    if not active_subs.data and not trial_subs.data:
                        await db.collection("customers").document(uid).set({
                            "status": "inactive", "plan": "none", "updated_at": firestore.SERVER_TIMESTAMP
                        }, merge=True)
                        log.info(f"❌ Webhook: Suscripción terminada/fallida para {uid}. Acceso revocado.")
                    else:
                        log.info(f"⚠️ Webhook: Fallo de pago ignorado para {uid}. El usuario mantiene otra suscripción activa/trialing.")
                except Exception as e:
                    log.error(f"Error verificando suscripciones paralelas para {uid}: {e}")

        return {"status": "success"}

    except stripe.error.SignatureVerificationError as e:
        log.error(f"❌ Firma de Webhook inválida: {e}")
        return Response(content="Invalid signature", status_code=400)
    except Exception as e:
        error_msg = f"Error interno: {type(e).__name__} - {str(e)}"
        log.error(f"💥 Error crítico en Webhook: {error_msg}", exc_info=True)
        return Response(content=error_msg, status_code=500)

@app.post("/create-portal-session", tags=["Billing"])
async def create_portal_session(request: Request, current_user: Dict[str, Any] = Depends(get_current_user)):
    try:
        try:
            body = await request.json()
            return_url = body.get("return_url")
        except Exception: return_url = None
        if not return_url: return_url = "https://pida-ai.com/"
        user_email = current_user.get("email")
        customers = stripe.Customer.list(email=user_email, limit=1)
        if not customers.data: raise HTTPException(status_code=404, detail="No se encontró un cliente asociado a este correo en Stripe.")
        stripe_customer_id = customers.data[0].id
        session = stripe.billing_portal.Session.create(customer=stripe_customer_id, return_url=return_url)
        return {"url": session.url}
    except Exception as e:
        log.error(f"Error generando sesión del portal: {e}")
        raise HTTPException(status_code=500, detail=str(e))
